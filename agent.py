from fastapi import FastAPI, UploadFile,Form
import fitz
from dotenv import load_dotenv
import os
import re
import json
from uuid import uuid4
from datetime import datetime
from sqlalchemy.ext.asyncio import create_async_engine, AsyncSession
from sqlalchemy.orm import sessionmaker, declarative_base
from sqlalchemy import Column, String, Text, DateTime,Integer
from sqlalchemy.dialects.postgresql import UUID
import uuid
from slugify import slugify
from fastapi.staticfiles import StaticFiles
from fastapi import Request
from fastapi.responses import HTMLResponse, JSONResponse
from io import BytesIO
import httpx
import traceback
from docx import Document
from apscheduler.schedulers.asyncio import AsyncIOScheduler




# ------------------------- Baza danych -------------------------
load_dotenv()
DATABASE_URL = os.getenv('DATABASE_URL')
EXCEL_OUTPUT_PATH = "shortlist_summary.xlsx"
engine = create_async_engine(DATABASE_URL, echo=False)
async_session = sessionmaker(engine, class_=AsyncSession, expire_on_commit=False)
Base = declarative_base()

class JobCriteria(Base):
    __tablename__ = "job_criteria"
    id = Column(UUID(as_uuid=True), primary_key=True, default=uuid.uuid4)
    form_id = Column(String, nullable=False)
    extracted_text = Column(Text)
    uploaded_at = Column(DateTime, default=datetime.utcnow)
    company = Column(String, nullable=True)
    position = Column(String, nullable=True)
# ------------------------- Aplikacja FastAPI -------------------------
app = FastAPI()
app.mount("/static", StaticFiles(directory="static"), name="static")
scheduler = AsyncIOScheduler()
@app.on_event("startup")
async def startup_event():
    start_scheduler()
@app.post("/expire_now/")
async def expire_now():
    await expire_old_submissions(async_session, EXPIRATION_DAYS)
    return {"status": "expire job executed"}


@app.get("/", response_class=HTMLResponse)
async def home():
    return FileResponse("main_panel.html")

@app.get("/upload.html", response_class=HTMLResponse)
async def upload_page():
    return FileResponse("upload.html")

@app.get("/send.html", response_class=HTMLResponse)
async def send_page():
    return FileResponse("send.html")

@app.get("/summary.html", response_class=HTMLResponse)
async def summary_page():
    return FileResponse("summary.html")

@app.get("/advanced_options", response_class=HTMLResponse)
async def advanced_options_page():
    return FileResponse("advanced_options.html")



from openai import OpenAI
api_key = os.getenv("OPENAI_API_KEY")

if not api_key:
    raise RuntimeError("Nie ustawiono zmiennej środowiskowej OPENAI_API_KEY!")
client = OpenAI(api_key=api_key)
# ------------------------- Pomocnicze -------------------------
def extract_text(file: BytesIO, filename: str) -> str:
    file.seek(0)
    if filename.endswith(".docx"):
        doc = Document(file)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text_parts.append(cell_text)
        return "\n".join(text_parts)
    elif filename.endswith(".pdf"):
        doc = fitz.open(stream=file.read(), filetype="pdf")
        return "\n".join(page.get_text().strip() for page in doc)
    elif filename.endswith(".odt"):
        from odf.opendocument import load
        from odf.text import P
        textdoc = load(file)
        paragraphs = textdoc.getElementsByType(P)
        return "\n".join([p.firstChild.data for p in paragraphs if p.firstChild])

    else:
        raise HTTPException(status_code=400, detail="Unsupported file format")

def extract_company_and_position(text: str) -> tuple[str, str]:
    prompt = f"""
You will receive a job description text. Extract the **company name** and the **job position** from it.

Return only a valid JSON object in the format:
{{
  "company": "...",
  "position": "..."
}}

TEXT:
{text}
"""
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}]
    )
    content = response.choices[0].message.content.strip()
    try:
        json_match = re.search(r'\{[\s\S]*}', content)
        if json_match:
            result = json.loads(json_match.group())
            return result.get("company", "unknown"), result.get("position", "unknown")
        else:
            raise ValueError("No JSON found")
    except Exception as e:
        print("AI extraction failed:", e)
        return "unknown", "unknown"

# ------------------------- ENDPOINT 1: Upload kryteriów -------------------------
ALLOWED_MIME_TYPES = [
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # .docx
    "application/vnd.oasis.opendocument.text"  # .odt
]
MAX_FILE_SIZE = 3 * 1024 * 1024  # 3 MB

@app.post("/upload_criteria/")
async def upload_criteria(criteria_file: UploadFile = Form(...)):
    criteria_bytes = BytesIO(await criteria_file.read())

    def validate_file(file: UploadFile, file_bytes: BytesIO):
        if file.content_type not in ALLOWED_MIME_TYPES:
            raise HTTPException(400, f"Unsupported type: {file.content_type}")

        file_bytes.seek(0, os.SEEK_END)
        size = file_bytes.tell()
        file_bytes.seek(0)
        if size > MAX_FILE_SIZE:
            raise HTTPException(413, "File too large. Max 3 MB.")
    validate_file(criteria_file, criteria_bytes)
    criteria_bytes.name = criteria_file.filename
    criteria_text = extract_text(criteria_bytes, criteria_bytes.name)

    if not criteria_text.strip():
        raise HTTPException(status_code=400, detail="Criteria file is empty or unreadable")


    company, position = extract_company_and_position(criteria_text)
    company_part = slugify(company.split()[0]) if company != "unknown" else "unknown"
    position_initials = ''.join([w[0] for w in position.split() if w[0].isalpha()]).lower() or "pos"
    date_part = datetime.today().strftime("%y%m%d")
    short_hash = uuid4().hex[:4]
    form_id = f"{company_part}-{position_initials}-{date_part}-{short_hash}"

    criteria_id = uuid.uuid4()

    async with async_session() as session:
        new_criteria = JobCriteria(
            id=criteria_id,
            form_id=form_id,
            extracted_text=criteria_text,
            company=company,
            position=position
        )

        session.add(new_criteria)
        await session.commit()

    return {
        "criteria_id": str(criteria_id),
        "form_id": form_id,
        "company": company,
        "position": position
    }
@app.get("/job_criteria_list/")
async def job_criteria_list():
    async with async_session() as session:
        result = await session.execute(select(JobCriteria))
        records = result.scalars().all()
        return [
            {
                "form_id": jc.form_id,
                "company": jc.company or "unknown",
                "position": jc.position or "unknown"
            }
            for jc in records
        ]
from fastapi import Query, Path

@app.delete("/delete_job_criteria")
async def delete_job_criteria(form_id: str = Query(...)):
    async with async_session() as session:
        result = await session.execute(select(JobCriteria).where(JobCriteria.form_id == form_id))
        criteria = result.scalar()
        if not criteria:
            raise HTTPException(status_code=404, detail="JobCriteria not found")
        await session.delete(criteria)
        await session.commit()
    return {"status": "deleted", "form_id": form_id}



from fastapi import Form, HTTPException, BackgroundTasks
from sqlalchemy import select
import os
import uuid


ZAPIER_WEBHOOK_URL = os.getenv('ZAPIER_WEBHOOK_URL')

@app.post("/send_to_zapier/")
async def send_to_zapier(
    form_id: str = Form(...),
    email: str = Form(...),
    background_tasks: BackgroundTasks = None
):
    async with async_session() as session:
        result = await session.execute(
            select(JobCriteria).where(JobCriteria.form_id == form_id)
        )
        exists = result.scalar()
        if not exists:
            raise HTTPException(status_code=404, detail="form_id not found in database")

    template_path = "templates/answers.docx"
    if not os.path.exists(template_path):
        raise HTTPException(status_code=404, detail="Template not found")

    doc = Document(template_path)
    from docx.shared import RGBColor

    para = doc.add_paragraph()
    run = para.add_run(f"FORM_ID::{form_id}")
    run.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()

    from slugify import slugify
    company_safe = slugify(exists.company or "unknown")
    position_safe = slugify(exists.position or "position")
    tmp_filename = f"form-{company_safe}_{position_safe}_{uuid.uuid4()}.docx"
    doc.save(tmp_filename)

    # Zapisz wysłanie do bazy
    async with async_session() as session:
        submission = CandidateSubmission(
            email=email,
            form_id=form_id,
            send_time=datetime.utcnow(),
            status="sent"
        )
        session.add(submission)
        await session.commit()

    background_tasks.add_task(send_to_zapier_background, tmp_filename, form_id, email)
    return {"status": "sent", "email": email, "form_id": form_id}

async def send_to_zapier_background(file_path: str, form_id: str, email: str):
    """
    Funkcja wysyłająca dane do Zapiera w tle.
    """
    try:
        with open(file_path, "rb") as file:
            files = {
                "file": (os.path.basename(file_path), file, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            }
            data = {"email": email}

            async with httpx.AsyncClient() as client:
                response = await client.post(ZAPIER_WEBHOOK_URL, data=data, files=files)
                print(f"Zapier response status: {response.status_code}, body: {response.text}")

            if response.status_code != 200:
                raise HTTPException(status_code=500, detail=f"Zapier error: {response.text}")

            print(f"Form sent to {email} via Zapier")
    except Exception as e:
        print(f"Failed to send form to Zapier: {str(e)}")
    finally:
        try:
            os.remove(file_path)
        except Exception as cleanup_error:
            print(f"Could not delete temporary file: {cleanup_error}")

# ------------------------- Odebranie danych -------------------------
class CandidateSubmission(Base):
    __tablename__ = "candidate_submissions"

    id = Column(Integer, primary_key=True, autoincrement=True)
    email = Column(String, nullable=False)
    form_id = Column(String, nullable=True)
    answers_json = Column(Text)
    submitted_at = Column(DateTime, default=datetime.utcnow)
    verdict = Column(String, nullable=True)
    send_time = Column(DateTime(timezone=True), nullable=True)
    status = Column(String, nullable=True)

# ------------------------- Odebranie danych -------------------------
@app.post("/receive_from_zapier/")

async def receive_from_zapier(request: Request):
    try:
        form = await request.form()
        email = form.get("email")
        file_url = form.get("file")

        if not email or not file_url:
            raise HTTPException(status_code=400, detail="Missing email or file URL")

        async with httpx.AsyncClient() as client:
            response = await client.get(file_url)
            response.raise_for_status()
            file_content = response.content

        doc_buf = BytesIO(file_content)
        document = Document(doc_buf)

        form_id = None
        for para in document.paragraphs:
            if para.text.strip().startswith("FORM_ID::"):
                form_id = para.text.strip().replace("FORM_ID::", "").strip()
                break

        if not form_id:
            raise HTTPException(status_code=400, detail="form_id missing in document")

        parsed_answers = {}
        try:
            doc_buf = BytesIO(file_content)
            document = Document(doc_buf)

            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if len(cells) >= 2 and cells[0]:
                        question = cells[0]
                        answer = cells[1]
                        parsed_answers[question] = answer
        except Exception as e:
            print("Could not extract text from document:", e)

        async with async_session() as session:
            existing_result = await session.execute(
                select(CandidateSubmission).where(
                    CandidateSubmission.email == email,
                    CandidateSubmission.form_id == form_id
                )
            )
            existing_submission = existing_result.scalar()

            if existing_submission:
                existing_submission.answers_json = json.dumps(parsed_answers)
                existing_submission.submitted_at = datetime.utcnow()
                existing_submission.verdict = None
                existing_submission.status = "received"
                session.add(existing_submission)
                await session.commit()
                candidate_submission = existing_submission
            else:
                new_submission = CandidateSubmission(
                    email=email,
                    form_id=form_id,
                    answers_json=json.dumps(parsed_answers)
                )
                session.add(new_submission)
                await session.commit()
                candidate_submission = new_submission

        await analyze_submission(form_id, candidate_submission.email)

        return JSONResponse({
            "status": "saved",
            "email": email,
            "form_id": form_id,
            "answers_count": len(parsed_answers)
        })

    except Exception as e:
        print("Exception occurred:")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Error processing Zapier data")

# ------------------------- Analiza na podstawie form_id -------------------------
async def analyze_submission(form_id: str, email: str):
    async with async_session() as session:
        job_result = await session.execute(select(JobCriteria).where(JobCriteria.form_id == form_id))
        job_criteria = job_result.scalar()

        candidate_result = await session.execute(
            select(CandidateSubmission).where(
                CandidateSubmission.form_id == form_id,
                CandidateSubmission.email == email
            )
        )

        candidate_submission = candidate_result.scalar()

        if not job_criteria or not candidate_submission:
            print(f"Brak danych dla form_id: {form_id}")
            return

        criteria_text = job_criteria.extracted_text
        answers = json.loads(candidate_submission.answers_json)
        prompt = f"""
        You are an expert recruitment assistant.

        You will receive:
        1. A job description with requirements and qualifications,
        2. A candidate's answers submitted via form.

        Your task is to analyze the candidate's fit to the job requirements only from sections: 'Expected knowledge', 'Qualifications' and 'Work experience' in criteria text.

Use **contextual understanding**. Even if the candidate doesn't use the same keywords, recognize equivalent experience or qualifications (e.g., "Worked at Ubisoft" implies gamedev).  
        Output one short sentence (max 10 words) that clearly states any **missing elements** from the candidate profile, such as:

        - missing formal qualifications
        - lack of relevant experience
        - lack of required certifications or skills
        - insufficient detail or vague answers

        If the candidate fits well and there are no gaps, you respond: "Strong match for the position".

        Do not list every requirement. Mention only **key gaps** (e.g. “Lacks formal qualifications and required certifications”).

        Return **only** the sentence. No formatting. No JSON. No explanations.

        ### Job Description:
        {criteria_text}

        ### Candidate's Answers:
        {json.dumps(answers, indent=2, ensure_ascii=False)}
        """

        response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}]
        )

        result = response.choices[0].message.content.strip()
        candidate_submission.verdict = result.strip()
        session.add(candidate_submission)
        await session.commit()
        if candidate_submission.answers_json and candidate_submission.verdict:
            await update_excel_summary(candidate_submission)
        else:
            print(f"Skipped Excel update for form_id={form_id}, missing answers or verdict.")

        print(f"Analyzing submission ID={candidate_submission.id}, email={candidate_submission.email}")

        print(f"\nWYNIK ANALIZY DLA form_id={form_id} ({datetime.now().isoformat()}):\n{result}\n")
from fastapi import FastAPI, HTTPException, Body
from sqlalchemy import update
from datetime import datetime, timezone, timedelta
from apscheduler.schedulers.asyncio import AsyncIOScheduler
from apscheduler.triggers.cron import CronTrigger
import json
from pathlib import Path
from fastapi import Body, HTTPException
from apscheduler.schedulers.asyncio import AsyncIOScheduler
from apscheduler.triggers.cron import CronTrigger
from fastapi.responses import JSONResponse
from datetime import datetime, timedelta, timezone
from sqlalchemy import update, select

CONFIG_PATH = Path("config.json")


# === Funkcje do obsługi pliku konfiguracyjnego ===
def load_expiration_days_from_file() -> int:
    if CONFIG_PATH.exists():
        try:
            with open(CONFIG_PATH, "r") as f:
                data = json.load(f)
                return int(data.get("expiration_days",0 ))
        except Exception:
            return 0
    else:
        return 0

def save_expiration_days_to_file(value: int):
    with open(CONFIG_PATH, "w") as f:
        json.dump({"expiration_days": value}, f)


# === Zmienna globalna ===
EXPIRATION_DAYS = load_expiration_days_from_file()


# === Endpointy API ===
@app.get("/get_expiration_time")
async def get_expiration_time():
    return {"expiration_time_days": EXPIRATION_DAYS}


@app.post("/update_expiration_time_global")
async def update_expiration_time_global(payload: dict = Body(...)):
    global EXPIRATION_DAYS
    expiration_days = payload.get("expiration_days")

    if not isinstance(expiration_days, int) or expiration_days <= 0:
        raise HTTPException(status_code=400, detail="Invalid expiration_days value")

    EXPIRATION_DAYS = expiration_days
    save_expiration_days_to_file(EXPIRATION_DAYS)

    return {"status": "ok", "value": EXPIRATION_DAYS}


# === Automatyczne wygaszanie zgłoszeń ===
async def expire_old_submissions(async_session, expire_limit_days: int):
    now_utc = datetime.now(timezone.utc).replace(tzinfo=None)
    expire_before = now_utc - timedelta(days=expire_limit_days)

    async with async_session() as session:
        # Pobieramy rekordy, które będą wygaszone
        result = await session.execute(
            select(CandidateSubmission)
            .where(CandidateSubmission.verdict.is_(None))
            .where(CandidateSubmission.send_time <= expire_before)
        )
        submissions_to_expire = result.scalars().all()

        # Aktualizujemy rekordy w bazie
        for submission in submissions_to_expire:
            submission.status = "expired"
            submission.verdict = "provided no form answers on time"
            submission.submitted_at = now_utc
            session.add(submission)

        await session.commit()

        # --- Aktualizacja Excela dla każdego wygaszonego rekordu ---
        for submission in submissions_to_expire:
            await update_excel_summary(submission)



async def expire_old_submissions_job():
    await expire_old_submissions(async_session, EXPIRATION_DAYS)


# === Harmonogram ===
scheduler = AsyncIOScheduler()

from zoneinfo import ZoneInfo

def start_scheduler():
    scheduler.add_job(
        expire_old_submissions_job,
        CronTrigger(hour=10, minute=13, timezone=ZoneInfo("Europe/Warsaw"))  # 10:00 czasu PL
    )
    scheduler.start()
    print("Scheduler started...")

import pandas as pd
from openpyxl.styles import Font, PatternFill
from openpyxl.utils import get_column_letter

async def update_excel_summary(submission: CandidateSubmission):
    file_exists = os.path.exists(EXCEL_OUTPUT_PATH)

    if file_exists:
        df = pd.read_excel(EXCEL_OUTPUT_PATH)
    else:
        df = pd.DataFrame(columns=["id", "email", "verdict", "resume", "shortlist_status"])

    new_row = {
        "email": submission.email,
        "verdict": submission.verdict,
        "resume": "yes",
        "shortlist_status": "yes" if submission.verdict and "trong match for the position" in submission.verdict.lower() else "no"
    }

    if not df.empty and "email" in df.columns:
        df.set_index("email", inplace=True)
        df.loc[submission.email] = new_row
        df.reset_index(inplace=True)
    else:
        df = pd.DataFrame([new_row])

    df["id"] = range(1, len(df) + 1)

    # Przestawiamy kolumny: id ma być pierwsze
    desired_order = ["id", "email", "verdict", "resume", "shortlist_status"]
    df = df[desired_order]

    with pd.ExcelWriter(EXCEL_OUTPUT_PATH, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name="Summary")
        worksheet = writer.sheets["Summary"]

        # Styl nagłówków
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="1D6F42", end_color="1D6F42", fill_type="solid")

        for col_num, column_title in enumerate(df.columns, 1):
            cell = worksheet.cell(row=1, column=col_num)
            cell.font = header_font
            cell.fill = header_fill

            column_letter = get_column_letter(col_num)
            max_length = max(
                [len(str(cell.value)) if cell.value else 0 for cell in worksheet[column_letter]]
            )
            adjusted_width = min(max_length + 2, 100)
            worksheet.column_dimensions[column_letter].width = adjusted_width



from fastapi.responses import FileResponse
from sqlalchemy import text
from fastapi.responses import StreamingResponse
import io
import pandas as pd
from openpyxl.styles import Font, PatternFill
from openpyxl.utils import get_column_letter

@app.get("/generate_summary_excel/")
async def generate_summary_excel():
    async with async_session() as session:
        result = await session.execute(select(CandidateSubmission))
        submissions = result.scalars().all()

    # Przygotowujemy słownik, gdzie dla każdego emaila zostawimy tylko najnowszy rekord z verdict
    latest_submissions = {}
    for s in submissions:
        if not s.verdict or s.verdict.strip() == "":
            continue  # pomijamy rekordy bez verdict
        if s.email not in latest_submissions:
            latest_submissions[s.email] = s
        else:
            if s.submitted_at > latest_submissions[s.email].submitted_at:
                latest_submissions[s.email] = s

    data = []
    for s in latest_submissions.values():
        data.append({
            "email": s.email,
            "verdict": s.verdict,
            "resume": "yes",
            "shortlist_status": "yes" if "trong match for the position" in s.verdict.lower() else "no"
        })

    df = pd.DataFrame(data)
    df["id"] = range(1, len(df) + 1)

    # Ustawiamy kolumnę 'id' jako pierwszą
    df = df[["id", "email", "verdict", "resume", "shortlist_status"]]

    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Summary")
        worksheet = writer.sheets["Summary"]

        # Styl nagłówków
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="1D6F42", end_color="1D6F42", fill_type="solid")

        for col_num, column_title in enumerate(df.columns, 1):
            cell = worksheet.cell(row=1, column=col_num)
            cell.font = header_font
            cell.fill = header_fill

            column_letter = get_column_letter(col_num)
            max_length = max(
                [len(str(c.value)) if c.value else 0 for c in worksheet[column_letter]]
            )
            adjusted_width = min(max_length + 2, 100)
            worksheet.column_dimensions[column_letter].width = adjusted_width

    output.seek(0)
    headers = {
        "Content-Disposition": "attachment; filename=summary.xlsx"
    }
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers=headers,
    )


@app.get("/summary_json/")
async def summary_json():
    file_path = "shortlist_summary.xlsx"
    if not os.path.exists(file_path):
        return []
    df = pd.read_excel(file_path)
    return df.to_dict(orient="records")
import shutil
from datetime import datetime

from fastapi import HTTPException
from sqlalchemy import select, delete

from sqlalchemy import delete
from sqlalchemy.ext.asyncio import AsyncSession

@app.post("/reset_summary/")
async def reset_summary():
    if os.path.exists(EXCEL_OUTPUT_PATH):
        os.remove(EXCEL_OUTPUT_PATH)

    async with async_session() as session:
        await session.execute(delete(CandidateSubmission).execution_options(synchronize_session=False))
        await session.commit()
        result = await session.execute(select(CandidateSubmission))
        records = result.scalars().all()
        print("Remaining records after delete:", records)
    return {"status": "reset"}

from fastapi import Query
from sqlalchemy import delete
from sqlalchemy.ext.asyncio import AsyncSession

@app.delete("/delete_summary_entry/")
async def delete_summary_entry(email: str = Query(...)):
    if not os.path.exists(EXCEL_OUTPUT_PATH):
        raise HTTPException(status_code=404, detail="Summary Excel not found")

    df = pd.read_excel(EXCEL_OUTPUT_PATH)
    original_len = len(df)
    df = df[df.email != email]

    if len(df) == original_len:
        raise HTTPException(status_code=404, detail="Email not found in summary")

    df.reset_index(drop=True, inplace=True)
    df["id"] = range(1, len(df) + 1)
    df.to_excel(EXCEL_OUTPUT_PATH, index=False)

    # --- USUNIĘCIE Z BAZY DANYCH ---
    async with async_session() as session:
        await session.execute(
            delete(CandidateSubmission).where(CandidateSubmission.email == email)
        )
        await session.commit()

    return {"status": "deleted", "email": email}


